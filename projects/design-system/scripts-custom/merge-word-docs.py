#!/usr/bin/env python3
"""
Merge multiple Word documents into a single document.
Excludes PROTECTED_FILES.docx and SETUP_PROTECTION.docx.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK

def merge_word_documents():
    # Define paths
    docs_dir = Path("documentation/wordfiles-temp")
    output_file = docs_dir / "MERGED_DESIGN_SYSTEM_DOCS.docx"
    
    # Files to exclude
    exclude_files = {"PROTECTED_FILES.docx", "SETUP_PROTECTION.docx"}
    
    if not docs_dir.exists():
        print(f"❌ Directory not found: {docs_dir}")
        return
    
    # Get all .docx files except excluded ones
    docx_files = [f for f in docs_dir.glob("*.docx") 
                  if f.name not in exclude_files and f.name != "MERGED_DESIGN_SYSTEM_DOCS.docx"]
    
    if not docx_files:
        print("❌ No .docx files found to merge")
        return
    
    # Sort files alphabetically for consistent order
    docx_files.sort(key=lambda x: x.name)
    
    print(f"📄 Found {len(docx_files)} files to merge:")
    for file in docx_files:
        print(f"   - {file.name}")
    
    # Create new document
    merged_doc = Document()
    
    # Add title page
    title = merged_doc.add_heading("Oblique Design System Documentation", 0)
    title_para = merged_doc.add_paragraph()
    title_para.add_run("Complete Documentation Package").bold = True
    merged_doc.add_paragraph(f"Generated: {Path.cwd().name}")
    merged_doc.add_paragraph()
    
    # Add table of contents placeholder
    merged_doc.add_heading("Table of Contents", 1)
    for i, file in enumerate(docx_files, 1):
        doc_name = file.stem.replace('-', ' ').replace('_', ' ').title()
        merged_doc.add_paragraph(f"{i}. {doc_name}")
    
    # Add page break before content
    merged_doc.add_page_break()
    
    # Merge each document
    for i, file_path in enumerate(docx_files):
        print(f"📖 Merging: {file_path.name}")
        
        try:
            # Load source document
            source_doc = Document(file_path)
            
            # Add document title
            doc_name = file_path.stem.replace('-', ' ').replace('_', ' ').title()
            merged_doc.add_heading(f"{i+1}. {doc_name}", 1)
            
            # Copy paragraphs properly
            for paragraph in source_doc.paragraphs:
                if paragraph.text.strip():  # Only copy non-empty paragraphs
                    new_para = merged_doc.add_paragraph()
                    
                    # Copy paragraph style and content
                    for run in paragraph.runs:
                        new_run = new_para.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        if run.font.size:
                            new_run.font.size = run.font.size
            
            # Copy tables
            for table in source_doc.tables:
                new_table = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                for i_row, row in enumerate(table.rows):
                    for i_col, cell in enumerate(row.cells):
                        new_table.cell(i_row, i_col).text = cell.text
            
            # Add page break between documents (except for last one)
            if i < len(docx_files) - 1:
                merged_doc.add_page_break()
                
        except Exception as e:
            print(f"⚠️  Warning: Could not merge {file_path.name}: {e}")
            continue
    
    # Save merged document
    try:
        merged_doc.save(output_file)
        print(f"\n✅ Successfully merged {len(docx_files)} documents")
        print(f"📁 Output file: {output_file}")
        print(f"📊 File size: {output_file.stat().st_size / 1024:.1f} KB")
        
        print(f"\n📋 Merged documents:")
        for file in docx_files:
            print(f"   ✅ {file.name}")
            
        print(f"\n🚫 Excluded files:")
        for excluded in exclude_files:
            print(f"   ❌ {excluded}")
            
    except Exception as e:
        print(f"❌ Error saving merged document: {e}")

if __name__ == "__main__":
    print("🔗 Merging Word documents...")
    print("=" * 60)
    merge_word_documents()
