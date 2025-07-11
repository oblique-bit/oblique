#!/usr/bin/env python3
"""
Generate Word Documents from Markdown Documentation
==================================================

Converts all markdown files in the documentation/ folder to Word (.docx) format
for offline reading and validation. Places output in documentation/wordfiles-temp/

Features:
- Professional footer with filename, generation date/time, and page numbers (Page X of Y)
- Proper formatting for headers, lists, code blocks, and basic markdown elements
- Clean, readable layout optimized for printing and sharing

Requirements:
- python-docx (install with: pip3 install python-docx markdown2)

Usage:
    python3 scripts-custom/generate-word-docs.py
"""

import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    import markdown2
    from datetime import datetime
except ImportError:
    print("❌ Missing required packages. Install with:")
    print("   pip3 install python-docx markdown2")
    sys.exit(1)

def add_footer_with_page_numbers(doc, filename):
    """Add footer with filename, date/time, and page numbers."""
    section = doc.sections[0]
    footer = section.footer
    
    # Get current date and time
    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    # Create footer paragraph with proper formatting
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add filename
    filename_run = footer_para.add_run(filename)
    filename_run.font.size = Pt(8)
    filename_run.font.name = 'Calibri'
    
    # Add separator and generation time
    separator_run = footer_para.add_run(f" | Generated: {current_time}")
    separator_run.font.size = Pt(8)
    separator_run.font.name = 'Calibri'
    
    # Add line break for page numbers
    footer_para.add_run("\n")
    
    # Add page numbers - center aligned
    page_para = footer.add_paragraph()
    page_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add "Page" text
    page_run = page_para.add_run("Page ")
    page_run.font.size = Pt(8)
    page_run.font.name = 'Calibri'
    
    # Add current page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    page_num_run = page_para.add_run("")
    page_num_run.font.size = Pt(8)
    page_num_run.font.name = 'Calibri'
    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)
    
    # Add " of " text
    of_run = page_para.add_run(" of ")
    of_run.font.size = Pt(8)
    of_run.font.name = 'Calibri'
    
    # Add total pages field
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = 'NUMPAGES'
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    
    total_run = page_para.add_run("")
    total_run.font.size = Pt(8)
    total_run.font.name = 'Calibri'
    total_run._r.append(fldChar3)
    total_run._r.append(instrText2)
    total_run._r.append(fldChar4)

def create_word_doc_from_markdown(md_content, title, filename):
    """Create a Word document from markdown content."""
    doc = Document()
    
    # Add title
    title_para = doc.add_heading(title, 0)
    
    # Add footer with filename, date/time, and page numbers
    add_footer_with_page_numbers(doc, filename)
    
    # Convert markdown to HTML then parse
    html_content = markdown2.markdown(md_content, extras=['fenced-code-blocks', 'tables'])
    
    # Simple conversion - split by lines and handle basic formatting
    lines = md_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Handle headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        # Handle code blocks
        elif line.startswith('```'):
            continue  # Skip code block markers
        elif line.startswith('    ') or line.startswith('\t'):
            # Code block content
            para = doc.add_paragraph(line, style='Code')
        # Handle lists
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. ') or any(line.startswith(f'{i}. ') for i in range(1, 10)):
            # Numbered list
            content = line.split('. ', 1)[1] if '. ' in line else line
            doc.add_paragraph(content, style='List Number')
        # Handle bold/italic (basic)
        else:
            # Regular paragraph
            if line:
                para = doc.add_paragraph()
                # Simple bold handling
                if '**' in line:
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            para.add_run(part)
                        else:
                            para.add_run(part).bold = True
                else:
                    para.add_run(line)
    
    return doc

def convert_md_to_docx(md_file, output_dir):
    """Convert a markdown file to Word document."""
    md_path = Path(md_file)
    docx_name = md_path.stem + '.docx'
    docx_path = output_dir / docx_name
    
    try:
        # Read markdown content
        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Create Word document
        doc = create_word_doc_from_markdown(md_content, md_path.stem.replace('-', ' ').title(), md_path.name)
        
        # Save document
        doc.save(str(docx_path))
        
        print(f"✅ Converted: {md_path.name} → {docx_name}")
        return True
        
    except Exception as e:
        print(f"❌ Failed to convert {md_path.name}: {e}")
        return False

def main():
    """Main function to convert all documentation markdown files to Word."""
    
    # Set up paths
    project_root = Path(__file__).parent.parent
    docs_dir = project_root / 'documentation'
    output_dir = docs_dir / 'wordfiles-temp'
    
    # Ensure output directory exists
    output_dir.mkdir(exist_ok=True)
    
    # Find all markdown files in documentation
    md_files = list(docs_dir.glob('*.md'))
    
    if not md_files:
        print("❌ No markdown files found in documentation/")
        return
    
    print(f"📄 Found {len(md_files)} markdown files in documentation/")
    print(f"📁 Output directory: {output_dir}")
    print("-" * 60)
    
    # Convert each markdown file
    success_count = 0
    for md_file in sorted(md_files):
        if convert_md_to_docx(md_file, output_dir):
            success_count += 1
    
    print("-" * 60)
    print(f"✅ Successfully converted {success_count}/{len(md_files)} files")
    print(f"📁 Word documents saved to: {output_dir}")
    print("\n💡 You can now:")
    print("   - Open the .docx files in Microsoft Word")
    print("   - Print them for offline reading")
    print("   - Share them with team members")
    
    # List generated files
    docx_files = list(output_dir.glob('*.docx'))
    if docx_files:
        print(f"\n📋 Generated files:")
        for docx_file in sorted(docx_files):
            print(f"   - {docx_file.name}")

if __name__ == '__main__':
    main()
