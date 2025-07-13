#!/usr/bin/env python3
"""
Convert Status Colors Documentation from Markdown to Word format
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def create_word_document():
    # Create new document
    doc = Document()
    
    # Set document title
    title = doc.add_heading('Status Colors Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introduction
    doc.add_heading('Overview', level=1)
    intro_para = doc.add_paragraph()
    intro_para.add_run('Status colors communicate system state, user feedback, and contextual information in the Oblique Design System. They have been redesigned for government web applications with clearer, more accessible naming that follows established signal color semantics.')
    
    # Semantic Status System
    doc.add_heading('Semantic Status System', level=2)
    semantic_para = doc.add_paragraph()
    semantic_para.add_run('Semantic statuses act as an overarching term and contain more than colors. Status colors complement status icons (TBD). For full accessibility where possible (e.g., in Pill component), we use:')
    
    # Add bullet points
    doc.add_paragraph('Status color', style='List Bullet')
    doc.add_paragraph('Status label', style='List Bullet')
    doc.add_paragraph('Status icon (out of scope in this doc, has to be created as doc and linked)', style='List Bullet')
    
    # Foundation
    doc.add_heading('Foundation', level=2)
    foundation_para = doc.add_paragraph()
    foundation_para.add_run('Primitive color values are inherited from the BK Design System for consistency across Swiss government applications. We\'ve enhanced these base primitive colors with semantic meaning and grouped them into two distinct classifications to provide clearer guidance for designers and developers.')
    
    # Status Reference Table
    doc.add_heading('Complete Status Reference Table', level=1)
    
    # Create table
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    
    # Table headers
    headers = ['Status', 'Type', 'Description', 'Components', 'Legacy Name', 'Inspiration', 'Change Status', 'Classification']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Make headers bold
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Table data
    statuses = [
        ['info', 'Informational', 'General information, tips, announcements', 'ob.c.infobox, ob.c.badge', '—', 'GOV.UK, USWDS', 'Unchanged', 'Reserved'],
        ['resolved', 'Success/Completion', 'Completed tasks, success messages', 'ob.c.badge, ob.c.pill', 'success', 'GitHub, Jira', 'Renamed', 'Reserved'],
        ['critical', 'Error/Urgent', 'System failures, urgent alerts', 'ob.c.badge, ob.c.infobox', 'error', 'Material, Atlassian', 'Renamed', 'Reserved'],
        ['attention', 'Warning/Caution', 'Warnings, review needed', 'ob.c.infobox, ob.c.badge', '—', 'USWDS, Atlassian', 'Renamed', 'Reserved'],
        ['pending', 'Workflow', 'Awaiting action, in queue', 'ob.c.pill, ob.c.badge', '—', 'Jira, GitHub', 'Added', 'Flexible'],
        ['confirmed', 'Approval', 'Verified, approved, confirmed', 'ob.c.pill, ob.c.badge', '—', 'GitHub, Atlassian', 'Added', 'Flexible'],
        ['progress', 'Processing', 'In progress, loading states', 'ob.c.pill, ob.c.badge', '—', 'Jira, GitHub', 'Added', 'Flexible'],
        ['scheduled', 'Planning', 'Scheduled, future items', 'ob.c.pill, ob.c.badge', '—', 'Jira, GitHub', 'Added', 'Flexible'],
        ['waiting', 'Queue', 'Waiting, queued, on hold', 'ob.c.pill, ob.c.badge', '—', 'Jira, GitHub', 'Added', 'Flexible'],
        ['closed', 'Archive', 'Closed, archived, ended', 'ob.c.pill, ob.c.badge', '—', 'GitHub', 'Added', 'Flexible'],
        ['disabled', 'Inactive', 'Disabled, unavailable', 'All ob.c components', '—', 'Material', 'Added', 'Flexible'],
        ['fatal', 'Emergency', 'Population danger, critical alerts', 'ob.c.infobox', '—', 'BK Design System', 'Added', 'Flexible']
    ]
    
    # Add data rows
    for status_data in statuses:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(status_data):
            row_cells[i].text = cell_data
            # Make status names bold
            if i == 0:
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    # Status Token Classification
    doc.add_heading('Status Token Classification System', level=1)
    classification_para = doc.add_paragraph()
    classification_para.add_run('Status tokens are classified as either Reserved or Flexible to provide clear guidance for designers and developers:')
    
    doc.add_paragraph('Reserved statuses (info, resolved, critical, attention) follow universal signal semantics and cannot be changed across systems', style='List Bullet')
    doc.add_paragraph('Flexible statuses (pending, confirmed, progress, etc.) are suggested for ticketing and workflow use, and may be remapped or extended by project designers', style='List Bullet')
    
    # Reserved Status Rules
    doc.add_heading('Reserved Status Rules', level=2)
    doc.add_paragraph('Status tokens: info, resolved, critical, attention')
    
    doc.add_heading('Characteristics:', level=3)
    doc.add_paragraph('Fixed across all projects - Cannot be renamed or remapped', style='List Bullet')
    doc.add_paragraph('Follow traffic light semantics - Based on universal signal colors', style='List Bullet')
    doc.add_paragraph('Cross-system consistency - Used across government applications', style='List Bullet')
    doc.add_paragraph('Color mapping locked - Primitive color references must not change', style='List Bullet')
    
    # Flexible Status Rules
    doc.add_heading('Flexible Status Rules', level=2)
    doc.add_paragraph('Status tokens: pending, confirmed, progress, scheduled, waiting, closed, disabled, fatal')
    
    doc.add_heading('Characteristics:', level=3)
    doc.add_paragraph('Project-adaptable - May be renamed or remapped by designers', style='List Bullet')
    doc.add_paragraph('Workflow-oriented - Designed for ticketing and task management', style='List Bullet')
    doc.add_paragraph('Extensible - Projects can add new flexible statuses', style='List Bullet')
    doc.add_paragraph('Reserved protection - Must not alter reserved statuses', style='List Bullet')
    
    # Reserved Status Naming Rationale
    doc.add_heading('Reserved Status Naming Rationale', level=1)
    rationale_para = doc.add_paragraph()
    rationale_para.add_run('The Reserved status tokens (info, resolved, critical, attention) were specifically redesigned for government web applications with enhanced semantic clarity. The renaming decisions from legacy terms follow established patterns in modern web application design:')
    
    # Critical
    doc.add_heading('critical (formerly "error")', level=2)
    doc.add_paragraph('Why "critical" is superior for web applications:')
    doc.add_paragraph('Broader semantic scope: "Critical" encompasses system failures, urgent alerts, and high-priority issues beyond simple form errors', style='List Bullet')
    doc.add_paragraph('Government context alignment: Government applications often deal with critical situations requiring immediate attention', style='List Bullet')
    doc.add_paragraph('Professional tone: "Critical" sounds more measured and professional for government communications', style='List Bullet')
    doc.add_paragraph('Industry standard: Major enterprise tools use "critical" for highest-priority issues', style='List Bullet')
    
    # Resolved
    doc.add_heading('resolved (formerly "success")', level=2)
    doc.add_paragraph('Why "resolved" is more application-focused:')
    doc.add_paragraph('Process-oriented language: "Resolved" indicates completion of a workflow or task', style='List Bullet')
    doc.add_paragraph('Task management semantics: Government applications often involve case management and workflow completion', style='List Bullet')
    doc.add_paragraph('Less celebratory tone: Appropriate for government contexts', style='List Bullet')
    doc.add_paragraph('Industry adoption: GitHub, Jira, and other enterprise platforms use "resolved"', style='List Bullet')
    
    # Technical Implementation
    doc.add_heading('Technical Implementation', level=1)
    
    doc.add_heading('Token Structure', level=2)
    doc.add_paragraph('ob.s.color.status.{status-name}.{property}.{contrast-level}.{inversity-variation}')
    
    doc.add_heading('Properties', level=3)
    doc.add_paragraph('fg - Foreground/text color for status content', style='List Bullet')
    doc.add_paragraph('bg - Background color for status containers', style='List Bullet')
    
    doc.add_heading('Contrast Levels', level=3)
    doc.add_paragraph('contrast-high - Maximum contrast for critical visibility', style='List Bullet')
    doc.add_paragraph('contrast-medium - Standard contrast for normal usage', style='List Bullet')
    doc.add_paragraph('contrast-low - Subtle contrast for secondary contexts', style='List Bullet')
    
    doc.add_heading('Inversity Variations', level=3)
    doc.add_paragraph('inversity-normal - Standard light theme', style='List Bullet')
    doc.add_paragraph('inversity-flipped - Dark theme / inverted contexts', style='List Bullet')
    
    # Usage Guidelines
    doc.add_heading('Usage Guidelines', level=1)
    usage_para = doc.add_paragraph()
    usage_para.add_run('The following examples demonstrate recommended implementation patterns. Actual developer implementations may vary based on project requirements, technical constraints, and team preferences.')
    
    # Accessibility
    doc.add_heading('Accessibility', level=1)
    
    doc.add_heading('Contrast Requirements', level=2)
    doc.add_paragraph('All status colors have been validated and meet WCAG 2.1 AA conformity standards for accessibility compliance.')
    
    doc.add_heading('Color Independence', level=2)
    doc.add_paragraph('Status meaning should never rely solely on color. Always pair status colors with:')
    doc.add_paragraph('Clear text labels', style='List Bullet')
    doc.add_paragraph('Icons or symbols', style='List Bullet')
    doc.add_paragraph('Contextual information', style='List Bullet')
    doc.add_paragraph('Appropriate ARIA labels', style='List Bullet')
    
    # Implementation Guidelines
    doc.add_heading('Implementation Guidelines', level=1)
    
    doc.add_heading('For Design System Maintainers:', level=2)
    doc.add_paragraph('Never modify reserved statuses - Color mappings and names are locked', style='List Bullet')
    doc.add_paragraph('Document flexible adaptations - Track project-specific modifications', style='List Bullet')
    doc.add_paragraph('Validate token updates - Ensure reserved statuses remain unchanged', style='List Bullet')
    doc.add_paragraph('Review new statuses - Classify as reserved or flexible', style='List Bullet')
    
    doc.add_heading('For Project Designers:', level=2)
    doc.add_paragraph('Use reserved statuses as-is - No modifications allowed', style='List Bullet')
    doc.add_paragraph('Adapt flexible statuses freely - Rename or remap as needed', style='List Bullet')
    doc.add_paragraph('Extend with new statuses - Add project-specific statuses as flexible', style='List Bullet')
    doc.add_paragraph('Document customizations - Track deviations from base system', style='List Bullet')
    
    doc.add_heading('For Developers:', level=2)
    doc.add_paragraph('Reference semantic tokens - Never use primitive colors directly', style='List Bullet')
    doc.add_paragraph('Follow component guidelines - Use appropriate status types per component', style='List Bullet')
    doc.add_paragraph('Test theme switching - Ensure inversity variations work correctly', style='List Bullet')
    
    # Note
    note_para = doc.add_paragraph()
    note_run = note_para.add_run('Note: ')
    note_run.bold = True
    note_para.add_run('All status colors automatically meet WCAG 2.1 AA contrast requirements. Accessibility validation is built into the design system tokens, so developers don\'t need to manually check contrast ratios.')
    
    return doc

if __name__ == "__main__":
    # Create the document
    document = create_word_document()
    
    # Save the document
    output_path = "/Users/davorradisic/vc git repo bit/oblique/projects/design-system/documentation/design-tokens/colors/colors-semantic-status.docx"
    document.save(output_path)
    print(f"Word document created successfully: {output_path}")
